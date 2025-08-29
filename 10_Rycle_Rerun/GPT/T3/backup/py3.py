import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_gap_images_and_csvs(output_dir):
    """
    Returns lists of GAP-highlighted PNGs and corresponding CSV files.
    """
    pngs = []
    csvs = []
    for fname in os.listdir(output_dir):
        if fname.startswith("Poly_") and fname.endswith("_gap.png"):
            pngs.append(fname)
        elif fname.startswith("Poly_") and fname.endswith("_gap_analysis.csv"):
            csvs.append(fname)
    pngs.sort()
    csvs.sort()
    return pngs, csvs

def generate_report(output_dir, report_path):
    # Gather outputs
    gap_pngs, gap_csvs = get_gap_images_and_csvs(output_dir)

    # Create document
    doc = Document()

    # Title
    doc.add_heading("Simulation Report on CLAHE and GAP Pixel Detection in Polymeric Images", 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents a simulation and analysis pipeline for detecting GAP pixels in polymeric images. "
        "The pipeline employs CLAHE enhancement and pixel-wise evaluation based on grayscale values and neighborhood criteria. "
        "Results, including CSV pixel data and GAP-highlighted images, are generated and summarized herein."
    )

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Image analysis is crucial in materials science and engineering, especially for studying polymeric structures. "
        "Accurate enhancement and characterization of features within such images facilitate further scientific understanding. "
        "This project focuses on the automated detection of 'GAP' pixels—pixels of interest defined by specific grayscale thresholds and neighborhood criteria—in images prefixed with 'Poly_'. "
        "The goal is to enhance these images, identify GAP pixels, and present the results in both tabular and visual formats."
    )

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "The analysis pipeline consists of several stages:\n"
        "1. **Image Selection & Enhancement:** All images with filenames beginning with 'Poly_' and in PNG or JPG format were selected from the specified input directory. "
        "Each image was processed using CLAHE (Contrast Limited Adaptive Histogram Equalization) with clipLimit=3 and tileGridSize=(10, 10) to improve local contrast and feature visibility.\n"
        "2. **Grayscale Conversion & GAP Detection:** The enhanced images were converted to grayscale. For each pixel, a 'GAP' flag was assigned. "
        "A pixel is considered a GAP pixel if its grayscale value lies between 1 and 150 (inclusive), and in at least one of the four main directions (up, down, left, right), "
        "there exist 25 contiguous neighboring pixels meeting the same grayscale condition.\n"
        "3. **Data Output:** For each image, two outputs were produced:\n"
        "   - A CSV file listing every pixel's coordinates, grayscale value, and GAP flag.\n"
        "   - A new binary PNG image where GAP pixels are shown in black (RGB 0,0,0) and all others in white (RGB 255,255,255).\n"
        "Automation and processing were implemented in Python, leveraging OpenCV for image enhancement, Pillow for grayscale conversion, and standard Python libraries for data management."
    )

    # Results
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "The pipeline processed all eligible images, generating detailed CSV tables and visual binary maps of GAP pixel locations. "
        "Below are the GAP-highlighted images for each processed input, demonstrating the effectiveness of the enhancement and detection algorithms."
    )

    # Insert images
    for png in gap_pngs:
        doc.add_paragraph(png, style='Intense Quote')
        img_path = os.path.join(output_dir, png)
        try:
            doc.add_picture(img_path, width=Inches(4))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Error displaying image {png}: {e}")

    # Save report
    doc.save(report_path)
    print(f"Word simulation report generated: {report_path}")

if __name__ == "__main__":
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup"
    report_path = os.path.join(output_dir, "GAP_Analysis_Simulation_Report.docx")
    generate_report(output_dir, report_path)
