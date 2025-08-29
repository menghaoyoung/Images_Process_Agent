import os
from docx import Document
from docx.shared import Inches
from datetime import datetime

def get_gap_image_files(output_dir, prefix="Poly_", suffix="_GAP_visual.png"):
    """Return a list of all GAP result images in the output directory."""
    return [f for f in os.listdir(output_dir) if f.startswith(prefix) and f.endswith(suffix)]

def add_images_to_doc(doc, image_paths):
    """Add all images to the Word document, each with a caption."""
    for img_path in image_paths:
        doc.add_picture(img_path, width=Inches(4))
        img_name = os.path.basename(img_path)
        doc.add_paragraph(f"Figure: {img_name}")
        doc.add_paragraph("")

def write_simulation_report(docx_path, output_dir):
    """
    Create a Word document report about the GAP detection simulation.
    The document contains Abstract, Introduction, Methods, and Results sections,
    and includes all GAP visualization images.
    """
    # Find all result images
    image_files = get_gap_image_files(output_dir)
    image_paths = [os.path.join(output_dir, f) for f in image_files]
    num_images = len(image_paths)
    now = datetime.now().strftime('%Y-%m-%d %H:%M')

    # Create doc
    doc = Document()
    doc.add_heading('Simulation Report: Automated GAP Pixel Detection in Polymer Images', 0)
    doc.add_paragraph(f'Report generated: {now}')
    doc.add_paragraph('')

    # Abstract
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This report presents a detailed simulation of an automated method for detecting and visualizing GAP pixels "
        "in a set of polymer microscopy images. Using advanced image enhancement and pixel-wise analysis algorithms, "
        "the workflow identifies regions of interest based on specific grayscale intensity and spatial criteria. "
        "Results are summarized with both tabular data and visual outputs for each image in the dataset."
    )

    # Introduction
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        "The accurate identification of structural features in polymer images is essential for understanding material "
        "properties and performance. Manual analysis of microscopy images is labor-intensive and prone to subjectivity. "
        "Automated image processing techniques, particularly those utilizing contrast enhancement and pixel-wise computation, "
        "offer a robust solution for the detection of critical features, such as GAP pixels. GAP pixels are defined as those "
        "with grayscale values in a specific range and with a contiguous spatial context, making their identification suitable "
        "for algorithmic approaches. This simulation aims to demonstrate an end-to-end pipeline for GAP detection and visualization."
    )

    # Methods
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        "The input dataset comprised all images with the 'Poly_' prefix in the source directory. Each image was first processed "
        "using Contrast Limited Adaptive Histogram Equalization (CLAHE) with a clip limit of 3 and a tile grid size of (10,10), "
        "which improved local contrast and highlighted subtle features. Enhanced images were then converted to grayscale, and each "
        "pixel was analyzed. A pixel was flagged as a GAP pixel if its grayscale value was between 1 and 150 (inclusive) and at least "
        "one of its four cardinal directions (up, down, left, right) contained 25 contiguous pixels also meeting the grayscale condition. "
        "For each image, a CSV file with per-pixel data (coordinates, grayscale value, GAP flag) was generated, and a new binary "
        "visualization image was created, with GAP pixels shown in black and non-GAP pixels in white. All processing steps were automated "
        "in Python, utilizing OpenCV for enhancement and Pillow for image handling."
    )

    # Results
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        f"A total of {num_images} polymer images were processed. For each, a corresponding GAP visualization was generated. "
        "The binary output images clearly differentiate GAP and non-GAP regions, providing an immediate visual representation "
        "of the spatial distribution of the detected pixels. The following figures display the GAP visualization images for all processed samples."
    )
    doc.add_paragraph("")

    # Insert images
    add_images_to_doc(doc, image_paths)

    doc.add_page_break()
    doc.save(docx_path)
    print(f"Simulation report written successfully: {docx_path}")

if __name__ == "__main__":
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup1"
    report_path = os.path.join(output_directory, 'Simulation_Report_GAP_Detection.docx')
    write_simulation_report(report_path, output_directory)
