import os
from docx import Document
from docx.shared import Inches
from datetime import datetime

# Define input and output directories (must match previous programs)
OUTPUT_DIR = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\GPT\T3\backup2"
REPORT_PATH = os.path.join(OUTPUT_DIR, "GAP_Simulation_Report.docx")

def get_image_and_csv_files(output_dir):
    """
    Find all _GAP_flag.png and _gap_analysis.csv files in the output directory.
    Returns a list of tuples: (basename, gap_img_path, csv_path)
    """
    files = os.listdir(output_dir)
    img_files = [f for f in files if f.endswith("_GAP_flag.png")]
    csv_files = [f for f in files if f.endswith("_gap_analysis.csv")]
    result = []
    for imgf in img_files:
        base = imgf.replace("_GAP_flag.png", "")
        csvf = f"{base}_gap_analysis.csv"
        if csvf in csv_files:
            result.append((base, os.path.join(output_dir, imgf), os.path.join(output_dir, csvf)))
    # Sort by base name
    result.sort()
    return result

def count_gap_stats(csv_path):
    """
    Read the CSV file and return total pixels and number with GAP_flag=1.
    """
    total = 0
    gap = 0
    try:
        import csv
        with open(csv_path, 'r', newline='') as f:
            reader = csv.DictReader(f)
            for row in reader:
                total += 1
                if row.get("GAP_flag", "0") == "1":
                    gap += 1
    except Exception as e:
        print(f"Error reading {csv_path}: {e}")
    return total, gap

def write_report(image_csv_pairs, report_path):
    doc = Document()
    doc.add_heading("Simulation Report: Automated GAP Pixel Detection in Polymeric Images", 0)

    # Abstract
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This report presents the results of an automated simulation for detecting GAP pixels in polymeric images. "
        "The analysis leverages computer vision techniques to enhance image contrast and algorithmically identify "
        "regions of interest based on defined grayscale and adjacency criteria. The workflow, results, and summary "
        "of the findings are detailed herein."
    )

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "Accurate detection of microstructural gaps within polymeric materials is vital for understanding material properties, "
        "performance, and failure mechanisms. Manual inspection is labor-intensive and subject to observer bias. "
        "This simulation utilizes digital image processing to automate the identification of GAP pixels in microscopy images "
        "with the 'Poly_' prefix. The approach aims to improve reproducibility, efficiency, and enable large-scale quantitative analysis."
    )

    # Methods
    doc.add_heading("Methods", level=1)
    doc.add_paragraph(
        "Images were sourced from the specified directory, filtered by the 'Poly_' prefix and supporting PNG/JPG formats. "
        "Each image underwent contrast enhancement using CLAHE (Contrast Limited Adaptive Histogram Equalization) with "
        "clipLimit=3 and tileGridSize=(10, 10) via OpenCV. The enhanced images were then converted to grayscale using Pillow (PIL). "
        "Each pixel was analyzed: a pixel was flagged as a GAP if its grayscale value ranged from 1 to 150 (inclusive) and if at least "
        "one adjacent direction (up, down, left, or right) contained 25 contiguous pixels also meeting the grayscale threshold. "
        "For each image, two outputs were generated: a CSV file containing the coordinates, grayscale value, and GAP flag of every pixel; "
        "and a PNG image highlighting GAP pixels in black and non-GAP pixels in white. The workflow was fully automated in Python."
    )

    # Results
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "The simulation processed all available polymeric images and generated GAP detection outputs. The following table summarizes "
        "the pixel-wise analysis for each image, including the total number of pixels and the count of detected GAP pixels. "
        "Representative GAP-flagged images are also presented below."
    )

    # Add a table with summary statistics
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Image Name'
    hdr_cells[1].text = 'Total Pixels'
    hdr_cells[2].text = 'GAP Pixels'
    hdr_cells[3].text = 'GAP Ratio (%)'

    for base, img_path, csv_path in image_csv_pairs:
        total, gap = count_gap_stats(csv_path)
        ratio = 100.0 * gap / total if total else 0.0
        row_cells = table.add_row().cells
        row_cells[0].text = base
        row_cells[1].text = str(total)
        row_cells[2].text = str(gap)
        row_cells[3].text = f"{ratio:.2f}"

    doc.add_paragraph("")

    # Add all GAP-flagged images
    doc.add_paragraph("GAP Detection Results (GAP=black, non-GAP=white):")
    for base, img_path, csv_path in image_csv_pairs:
        doc.add_paragraph(f"{base}_GAP_flag.png")
        try:
            doc.add_picture(img_path, width=Inches(3.5))
        except Exception as e:
            doc.add_paragraph(f"[Image could not be loaded: {e}]")

    # Add date and software note
    doc.add_paragraph(
        f"\nReport generated automatically on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} "
        "using Python 3, OpenCV, Pillow, and python-docx."
    )

    doc.save(report_path)
    print(f"Word report generated: {report_path}")

if __name__ == "__main__":
    pairs = get_image_and_csv_files(OUTPUT_DIR)
    if not pairs:
        print("No GAP analysis images or CSV files found in output directory.")
    else:
        write_report(pairs, REPORT_PATH)
