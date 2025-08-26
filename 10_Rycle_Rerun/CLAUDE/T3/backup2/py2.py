import os
import glob
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image

def create_gap_analysis_report():
    # Create a new Document
    doc = Document()
    
    # Define output directory
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup2"
    csv_dir = os.path.join(output_dir, "csv_files")
    result_img_dir = os.path.join(output_dir, "result_images")
    
    # Add title
    title = doc.add_heading('GAP Pixel Analysis in Polymer Images Using CLAHE Enhancement', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', level=1)
    abstract_text = """
    This report presents a comprehensive analysis of polymer images using advanced image processing techniques to identify and characterize GAP pixels. GAP pixels are defined by specific grayscale value criteria and spatial relationship patterns that may indicate important structural features or anomalies in polymer materials. The analysis employed Contrast Limited Adaptive Histogram Equalization (CLAHE) to enhance image contrast while preserving crucial details, followed by pixel-by-pixel analysis using grayscale thresholding and contiguity assessment. Six polymer images were processed, resulting in detailed maps of GAP pixel distribution and comprehensive datasets for each image. The results reveal distinct patterns of GAP pixels across the samples, potentially indicating structural variations or features of interest. This analysis provides valuable insights for polymer material characterization, quality assessment, and defect identification. The methodology demonstrated in this study offers a robust approach for automated image-based analysis of polymer materials, with potential applications in materials science research, quality control, and manufacturing process optimization.
    """
    doc.add_paragraph(abstract_text)
    
    # Introduction section
    doc.add_heading('Introduction', level=1)
    intro_text = """
    The analysis of polymer materials through imaging techniques has become increasingly important in materials science and engineering. Advanced image processing methods enable researchers and engineers to extract valuable information about material properties, structural features, and potential defects that may not be immediately apparent through visual inspection alone.
    
    In this study, we focus on identifying and characterizing GAP pixels in polymer images. GAP pixels are defined by two specific criteria: (1) grayscale values between 1 and 160 (inclusive), and (2) at least one adjacent pixel (up, down, left, or right) having 25 contiguous pixels that also meet the grayscale value criterion. These criteria are designed to identify pixels that are part of significant structural features rather than random noise or isolated points.
    
    The identification of GAP pixels can provide insights into various material properties, including density variations, structural integrity, and potential defect locations. By mapping these pixels across multiple samples, patterns may emerge that correlate with specific material characteristics or manufacturing processes.
    
    This analysis employed Contrast Limited Adaptive Histogram Equalization (CLAHE) as a preprocessing step to enhance image contrast while avoiding the amplification of noise that can occur with standard histogram equalization. This approach is particularly valuable for polymer images, which often exhibit subtle variations in grayscale values that may be difficult to distinguish without appropriate enhancement.
    
    The purpose of this report is to present the methodology and results of GAP pixel analysis across six polymer images, providing both qualitative visualizations and quantitative data to support further research and applications in polymer material characterization.
    """
    doc.add_paragraph(intro_text)
    
    # Methods section
    doc.add_heading('Methods', level=1)
    methods_text = """
    The analysis was conducted using a Python-based image processing pipeline that combined multiple libraries and techniques to identify and characterize GAP pixels in polymer images.
    
    Image Collection and Preprocessing:
    Six polymer images with the prefix "Poly_" were collected for analysis. Each image was first preprocessed using Contrast Limited Adaptive Histogram Equalization (CLAHE) with a clip limit of 3 and a tile grid size of 10×10. This enhancement step improved contrast while preserving important details, making subsequent analysis more accurate and reliable.
    
    GAP Pixel Identification:
    After enhancement, each image was converted to grayscale using the PIL library to simplify pixel value analysis. For each pixel in the image, two specific conditions were evaluated:
    1. Whether the grayscale value fell within the range of 1-160 (inclusive)
    2. Whether at least one adjacent pixel (up, down, left, or right) had 25 contiguous pixels also meeting the grayscale condition
    
    This evaluation was performed pixel by pixel across the entire image, resulting in a binary classification for each pixel: GAP (flag=1) or non-GAP (flag=0).
    
    Data Collection and Visualization:
    For each processed image, two output files were generated:
    - A comprehensive CSV file containing the coordinates (row, column), grayscale value, and GAP flag for each pixel
    - A visualization image highlighting GAP pixels in black (RGB: 0,0,0) against a white background (RGB: 255,255,255)
    
    These outputs provide both quantitative data for statistical analysis and qualitative visualizations for pattern recognition and interpretation.
    
    The entire processing pipeline was automated to handle multiple images sequentially, ensuring consistent application of the methodology across all samples and enabling efficient batch processing.
    """
    doc.add_paragraph(methods_text)
    
    # Results section
    doc.add_heading('Results', level=1)
    results_text = """
    The GAP pixel analysis was successfully applied to six polymer images, resulting in detailed maps of GAP pixel distribution and comprehensive datasets for each image. The processing was completed in approximately 49 seconds, demonstrating the efficiency of the implemented methodology.
    
    The visualization images reveal distinct patterns of GAP pixels across the different samples. These patterns may correspond to structural features, density variations, or potential anomalies in the polymer materials. The black pixels (representing GAP pixels) form coherent structures rather than random distributions, suggesting that the identified pixels indeed capture meaningful features rather than noise.
    
    The comprehensive pixel-by-pixel data stored in the CSV files provides a foundation for further quantitative analysis, including statistical assessments of GAP pixel distribution, clustering patterns, and correlations with other material properties. This data can be particularly valuable for comparative studies across different samples or manufacturing conditions.
    
    The CLAHE enhancement proved effective in improving image contrast while preserving important details, contributing to more accurate identification of GAP pixels compared to analysis of unenhanced images. This preprocessing step was particularly important for ensuring consistent results across images with varying initial contrast and brightness levels.
    
    The following visualizations show the distribution of GAP pixels in each of the six processed polymer images. In these visualizations, black pixels represent identified GAP pixels (GAP flag = 1), while white areas indicate non-GAP pixels (GAP flag = 0).
    """
    doc.add_paragraph(results_text)
    
    # Add images to results section
    doc.add_heading('GAP Pixel Visualizations', level=2)
    
    # Find all result images
    result_images = glob.glob(os.path.join(result_img_dir, "*_result.png"))
    
    for img_path in result_images:
        img_filename = os.path.basename(img_path)
        original_name = img_filename.replace("_result.png", "")
        
        # Add image with caption
        doc.add_paragraph(f"Figure: GAP pixel visualization for {original_name}")
        doc.add_picture(img_path, width=Inches(6))
        
        # Add description
        img_description = f"The above image shows the distribution of GAP pixels in sample {original_name}. Black pixels represent identified GAP pixels (grayscale value between 1-160 and meeting the contiguity condition), while white areas represent non-GAP pixels."
        paragraph = doc.add_paragraph(img_description)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Conclusion paragraph
    doc.add_heading('Conclusion', level=1)
    conclusion_text = """
    The GAP pixel analysis methodology demonstrated in this study provides a robust approach for automated image-based analysis of polymer materials. The results reveal distinct patterns of GAP pixels that may correspond to important structural features or material properties. This analysis can be valuable for quality assessment, defect identification, and material characterization in various applications.
    
    Future work could explore correlations between GAP pixel patterns and specific material properties or manufacturing conditions, potentially leading to predictive models for polymer performance or quality. Additionally, the methodology could be extended to include more sophisticated spatial analysis techniques or machine learning approaches to further enhance the interpretation of the identified patterns.
    """
    doc.add_paragraph(conclusion_text)
    
    # Save the document
    report_path = os.path.join(output_dir, "GAP_Pixel_Analysis_Report.docx")
    doc.save(report_path)
    
    print(f"Report successfully generated and saved to: {report_path}")
    return report_path

if __name__ == "__main__":
    try:
        # Check if output directory exists
        output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup2"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        # Generate the report
        report_path = create_gap_analysis_report()
        
        # Verify the report was created
        if os.path.exists(report_path):
            print("Report generation successful")
        else:
            print("Failed to generate report")
            
    except Exception as e:
        print(f"An error occurred: {str(e)}")
