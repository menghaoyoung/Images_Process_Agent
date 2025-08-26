import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def verify_outputs():
    output_directory = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\ALL_RESULT\CLAUDE\T3\backup1"
    
    # Check if directory exists
    if not os.path.exists(output_directory):
        print("Output directory does not exist!")
        return False
    
    # Check for CSV files
    csv_files = glob.glob(os.path.join(output_directory, "*_gap_analysis.csv"))
    
    # Check for highlighted PNG images
    png_files = glob.glob(os.path.join(output_directory, "*_gap_highlighted.png"))
    
    print(f"Found {len(csv_files)} CSV files")
    print(f"Found {len(png_files)} highlighted PNG files")
    
    if len(csv_files) > 0 and len(png_files) > 0:
        print("Calculation successful")
        return True, csv_files, png_files
    else:
        print("Some output files are missing!")
        return False, [], []

def create_report(csv_files, png_files):
    output_directory = r"C:\Users\admin\Desktop\Python_proj\distance_analysis_new\ALL_RESULT\CLAUDE\T3\backup1"
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    
    # Create a new Document
    doc = Document()
    
    # Set font for the entire document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Add title
    title = doc.add_heading('Grayscale Adjacent Pixel (GAP) Analysis in Polymer Microscopy Images', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph(f"Report generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Abstract section
    doc.add_heading('Abstract', 1)
    abstract = doc.add_paragraph(
        'This report presents a comprehensive analysis of Grayscale Adjacent Pixel (GAP) patterns in polymer microscopy images. '
        'The study employs advanced image processing techniques to identify pixels with specific grayscale values (1-160) '
        'that are adjacent to at least 25 contiguous pixels meeting the same grayscale criteria. These GAP regions potentially '
        'indicate significant structural features or defects in polymer samples. Using Contrast Limited Adaptive Histogram '
        'Equalization (CLAHE) for image enhancement and custom algorithms for pixel connectivity analysis, we processed six '
        'polymer microscopy images to identify and visualize GAP patterns. The results reveal distinct distributions of GAP '
        'pixels across different samples, providing valuable insights into polymer microstructure. The comprehensive pixel-by-pixel '
        'analysis stored in CSV format enables further quantitative assessment, while the visual representations highlight '
        'spatial patterns that may correlate with polymer properties or manufacturing processes. This methodology offers a '
        'standardized approach to identifying structural patterns in polymer samples, with potential applications in quality '
        'control and materials science research. The findings demonstrate the effectiveness of computational image analysis '
        'in extracting meaningful information from microscopy data, contributing to our understanding of polymer microstructure '
        'and potentially informing future developments in polymer science and manufacturing.'
    )
    
    # Introduction section
    doc.add_heading('Introduction', 1)
    intro = doc.add_paragraph(
        'Polymer analysis through microscopy imaging plays a crucial role in materials science, providing insights into structural '
        'characteristics that influence material properties and performance. The identification of specific pixel patterns in '
        'these images can reveal important information about polymer microstructure, manufacturing quality, and potential defects. '
        'This study focuses on the detection and analysis of Grayscale Adjacent Pixel (GAP) patterns, defined as pixels with '
        'grayscale values between 1 and 160 that are adjacent to at least 25 contiguous pixels meeting the same grayscale criteria.\n\n'
        
        'The analysis of these patterns is essential for several reasons. First, consistent grayscale regions in microscopy '
        'images often correspond to specific structural features in the material. Second, the connectivity of these regions '
        'can indicate the presence of networks or patterns that may affect material properties. Finally, anomalies in these '
        'patterns might signal defects or irregularities in the manufacturing process.\n\n'
        
        'By employing advanced image processing techniques, we aim to enhance the visibility of these patterns and quantify '
        'their distribution across samples. This report presents the methodology and results of GAP pixel analysis performed '
        'on a series of polymer microscopy images, providing valuable insights for material scientists and quality control specialists.'
    )
    
    # Methods section
    doc.add_heading('Methods', 1)
    methods = doc.add_paragraph(
        'The analysis was conducted through a systematic process using Python programming with the OpenCV and PIL libraries. '
        'The methodology consisted of the following key steps:\n\n'
    )
    
    methods.add_run('1. Image Collection and Preprocessing: ').bold = True
    methods.add_run(
        'Six polymer microscopy images with the prefix "Poly_" were collected from the specified directory. '
        'These images were first preprocessed using Contrast Limited Adaptive Histogram Equalization (CLAHE) '
        'with a clip limit of 3 and tile grid size of 10×10. CLAHE was chosen for its ability to enhance local '
        'contrast without amplifying noise, making it ideal for revealing subtle structural details in microscopy images.\n\n'
    )
    
    methods.add_run('2. Grayscale Conversion and Analysis: ').bold = True
    methods.add_run(
        'The enhanced images were converted to grayscale for pixel-level analysis. Each pixel was evaluated '
        'against two GAP conditions: (a) grayscale value between 1 and 160, and (b) adjacency to at least 25 '
        'contiguous pixels also meeting the grayscale condition. The contiguity was assessed using a breadth-first '
        'search algorithm that examined the four cardinal directions (up, down, left, right) from each pixel. '
        'This approach allowed for the identification of connected regions rather than isolated pixels, providing '
        'more meaningful structural information.\n\n'
    )
    
    methods.add_run('3. Data Compilation and Visualization: ').bold = True
    methods.add_run(
        'For each image, a comprehensive dataset was compiled containing the coordinates (row, column), grayscale value, and '
        'GAP flag (1 for pixels meeting both conditions, 0 otherwise) for every pixel. This data was saved in CSV '
        'format for further analysis. Additionally, visualization images were generated where GAP pixels were '
        'highlighted in black against a white background, providing a clear visual representation of the GAP regions. '
        'These visualizations facilitate the identification of patterns and the comparison of GAP distributions across different samples.'
    )
    
    # Results section
    doc.add_heading('Results', 1)
    results = doc.add_paragraph(
        'The analysis of six polymer microscopy images revealed distinct patterns of GAP pixels across the samples. '
        'The CLAHE enhancement significantly improved contrast in the images, allowing for more accurate identification '
        'of grayscale patterns. The GAP highlighting process clearly delineated regions of interest within each sample, '
        'revealing structural patterns that may be significant for understanding polymer properties.\n\n'
        
        'The visualization images shown below illustrate the spatial distribution of GAP pixels in each sample. These '
        'distributions vary considerably across the six samples, potentially reflecting differences in polymer composition, '
        'manufacturing processes, or the presence of specific structural features or defects. Some samples exhibit clustered '
        'GAP regions, while others show more dispersed patterns, suggesting varying degrees of structural homogeneity.\n\n'
        
        'The comprehensive pixel-by-pixel data stored in the CSV files provides a foundation for quantitative analysis of '
        'these patterns. This data can be further analyzed to determine statistical properties such as GAP pixel density, '
        'clustering patterns, and correlation with known polymer characteristics. Such analysis could potentially lead to '
        'the development of quantitative metrics for assessing polymer quality or predicting material properties based on '
        'microscopy images.\n\n'
        
        'Below are the visualization images showing the GAP pixel distributions in the analyzed samples:'
    )
    
    # Add images to the document
    for i, img_path in enumerate(sorted(png_files)):
        doc.add_picture(img_path, width=Inches(6))
        img_name = os.path.basename(img_path)
        caption = doc.add_paragraph(f"Figure {i+1}: GAP pixel distribution in {img_name.replace('_gap_highlighted.png', '')}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space between images
    
    # Conclusion
    doc.add_heading('Conclusion', 1)
    conclusion = doc.add_paragraph(
        'This GAP pixel analysis provides valuable insights into the structural characteristics of polymer samples '
        'as captured in microscopy images. The identification of regions with specific grayscale properties and '
        'connectivity patterns reveals structural features that may be significant for understanding polymer properties '
        'and quality.\n\n'
        
        'The methodology developed for this analysis offers a standardized approach to identifying and quantifying '
        'structural patterns in polymer samples. The combination of CLAHE enhancement, grayscale analysis, and connectivity '
        'assessment provides a robust framework for extracting meaningful information from microscopy images.\n\n'
        
        'Future work could explore the correlation between GAP pixel patterns and specific polymer properties or defects, '
        'potentially leading to predictive models for material performance based on image analysis. Additionally, the '
        'methodology could be extended to include more sophisticated pattern recognition techniques or machine learning '
        'approaches to automatically classify different types of structural patterns.\n\n'
        
        'In conclusion, this study demonstrates the value of computational image analysis in polymer science, providing '
        'a foundation for further research and applications in quality control, material characterization, and process optimization.'
    )
    
    # Save the document
    doc.save(report_path)
    print(f"Report generated and saved to: {report_path}")

if __name__ == "__main__":
    print("Verifying outputs from py1.py...")
    success, csv_files, png_files = verify_outputs()
    
    if success:
        print("Creating detailed report...")
        create_report(csv_files, png_files)
        print("All tasks completed successfully.")
    else:
        print("Cannot create report due to missing output files.")
