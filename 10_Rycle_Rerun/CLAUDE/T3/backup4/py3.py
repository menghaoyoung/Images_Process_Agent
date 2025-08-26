# py3.py - Generate a detailed report in Word format
import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import io
import datetime

def create_gap_analysis_report():
    # Define paths
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup4"
    result_image_directory = os.path.join(output_directory, "result_images")
    enhanced_directory = os.path.join(output_directory, "enhanced")
    csv_directory = os.path.join(output_directory, "csv")
    report_path = os.path.join(output_directory, "GAP_Analysis_Report.docx")
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Grayscale Adjacent Pixel (GAP) Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
    
    # Add Abstract section
    doc.add_heading('Abstract', 1)
    abstract = doc.add_paragraph()
    abstract.add_run(
        'This report presents a comprehensive analysis of images using the Grayscale Adjacent Pixel (GAP) methodology. '
        'The study examined five polygon-based images, applying advanced image processing techniques to identify '
        'pixels with specific grayscale characteristics and spatial relationships. Using Contrast Limited Adaptive '
        'Histogram Equalization (CLAHE) for image enhancement, the analysis focused on pixels with grayscale values '
        'between 1 and 150 that are part of or adjacent to extended structures with similar intensity profiles. '
        'The GAP approach defined significant pixels as those meeting two specific criteria: having a grayscale value '
        'within the specified range and having at least one adjacent direction (up, down, left, or right) with 25 '
        'contiguous pixels also meeting the grayscale condition. The results, visualized through binary classification '
        'images, reveal distinctive patterns within the processed images that may not be immediately apparent through '
        'visual inspection alone. These findings demonstrate the effectiveness of the GAP methodology in extracting '
        'meaningful structural information from grayscale imagery, with potential applications in material science, '
        'medical imaging, pattern recognition, and image segmentation. The comprehensive pixel-by-pixel data, stored '
        'in corresponding CSV files, enables further quantitative analysis beyond the visual representation provided '
        'in this report.'
    )
    
    # Add Introduction section
    doc.add_heading('Introduction', 1)
    intro = doc.add_paragraph()
    intro.add_run(
        'Image analysis plays a crucial role in extracting meaningful information from visual data across various '
        'scientific and industrial applications. This report focuses on a specific methodology termed Grayscale '
        'Adjacent Pixel (GAP) analysis, which examines both the intensity values of individual pixels and their '
        'spatial relationships with neighboring pixels.\n\n'
        'The purpose of this analysis is to identify and highlight structures within images that meet specific '
        'criteria related to grayscale values and contiguity. By focusing on pixels with grayscale values between '
        '1 and 150 (inclusive) that have at least one adjacent pixel (up, down, left, or right) with 25 contiguous '
        'pixels meeting the same grayscale condition, we aim to extract meaningful patterns that may represent '
        'important features or structures within the images.\n\n'
        'This approach differs from traditional thresholding or edge detection methods by considering not just the '
        'intensity of individual pixels but also their relationship to extended structures within the image. This '
        'makes the GAP methodology particularly valuable for identifying features that are characterized by both their '
        'intensity values and their spatial extent, providing insights that might be missed by simpler image analysis '
        'techniques.'
    )
    
    # Add Methods section
    doc.add_heading('Methods', 1)
    methods = doc.add_paragraph()
    methods.add_run(
        'The analysis was conducted using a systematic computational approach implemented in Python, leveraging '
        'several key libraries including OpenCV (cv2), PIL (Pillow), and NumPy. The methodology consisted of the '
        'following key steps:\n\n'
        '1. Image Collection and Preprocessing: Five images with the "Poly_" prefix were collected from the specified '
        'directory. These images were then enhanced using Contrast Limited Adaptive Histogram Equalization (CLAHE) '
        'with a clip limit of 3 and a tile grid size of 10×10. CLAHE was selected as the enhancement technique to '
        'improve contrast and feature visibility while avoiding the noise amplification that can occur with standard '
        'histogram equalization.\n\n'
        '2. Grayscale Conversion: The enhanced images were converted to grayscale using the PIL library to focus '
        'the analysis on intensity values rather than color information.\n\n'
        '3. GAP Analysis: Each pixel in the grayscale images was analyzed according to two specific criteria:\n'
        '   a. The pixel must have a grayscale value between 1 and 150 (inclusive).\n'
        '   b. At least one adjacent pixel (up, down, left, or right) must have 25 contiguous pixels that also meet '
        'the grayscale condition.\n\n'
        '4. Data Recording: For each image, a comprehensive CSV file was generated containing the coordinates '
        '(row, column), grayscale value, and GAP flag (0 or 1) for each pixel. This provides a detailed record of '
        'the analysis results for further quantitative assessment.\n\n'
        '5. Visualization: Binary images were created to visualize the GAP analysis results, with black pixels '
        '(RGB: 0,0,0) representing pixels that met the GAP criteria (flag=1) and white pixels (RGB: 255,255,255) '
        'representing those that did not (flag=0). This binary representation provides a clear visualization of the '
        'structures identified by the GAP analysis.\n\n'
        'The entire process was automated through a Python script, which efficiently processed all five images. '
        'This automation demonstrates the scalability of the approach for larger datasets and more complex analysis '
        'requirements.'
    )
    
    # Add Results section
    doc.add_heading('Results', 1)
    results = doc.add_paragraph()
    results.add_run(
        'The GAP analysis successfully identified and highlighted structures within the five processed images. '
        'The binary classification provided a clear visualization of features that met the specific GAP criteria, '
        'revealing patterns that may not be immediately apparent in the original images.\n\n'
        'The results for each image are presented below, showing the binary classification of pixels according to '
        'the GAP criteria. Black regions represent pixels that met both conditions: having grayscale values between '
        '1 and 150, and being part of or adjacent to extended structures of similar intensity.\n\n'
    )
    
    # Add images to the results section
    result_images = sorted(glob.glob(os.path.join(result_image_directory, '*_gap_result.png')))
    
    for img_path in result_images:
        base_name = os.path.basename(img_path)
        original_name = base_name.replace('_gap_result.png', '')
        
        # Add subheading for each image
        doc.add_heading(f"Analysis Results for {original_name}", 2)
        
        try:
            # Open image with PIL
            with Image.open(img_path) as img:
                # Resize if needed (keeping aspect ratio)
                width, height = img.size
                max_width = 6  # inches
                if width > max_width * 96:  # 96 dpi is common for Word
                    ratio = max_width * 96 / width
                    new_width = int(width * ratio)
                    new_height = int(height * ratio)
                    img = img.resize((new_width, new_height))
                
                # Save to memory buffer
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr.seek(0)
                
                # Add to document
                doc.add_picture(img_byte_arr, width=Inches(6))
                
                # Add caption
                caption = doc.add_paragraph(f"Figure: GAP analysis result for {original_name}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add description
                doc.add_paragraph(
                    f"The image above shows the GAP analysis result for {original_name}. Black pixels represent "
                    "areas that meet the GAP criteria (grayscale value between 1-150 and having at least one "
                    "adjacent direction with 25 contiguous pixels meeting the grayscale condition), while white "
                    "pixels represent areas that do not meet these criteria."
                )
                
        except Exception as e:
            doc.add_paragraph(f"Error adding image: {str(e)}")
    
    # Add conclusion to results
    doc.add_paragraph()
    conclusion = doc.add_paragraph()
    conclusion.add_run(
        'The binary visualizations clearly show the distribution of pixels meeting the GAP criteria across all five '
        'images. These results demonstrate the effectiveness of the GAP methodology in identifying extended structures '
        'within grayscale imagery. The comprehensive pixel data stored in the corresponding CSV files enables further '
        'quantitative analysis beyond the visual representation.\n\n'
        'The patterns revealed through this analysis could be valuable for various applications, including feature '
        'extraction, pattern recognition, and image segmentation. The methodology\'s consideration of both pixel '
        'intensity and spatial relationships provides a more nuanced understanding of the structures present in the '
        'images compared to simple thresholding techniques.\n\n'
        'Future work could explore variations in the GAP criteria, such as adjusting the grayscale value range or '
        'the number of contiguous pixels required, to optimize the analysis for specific applications or image types. '
        'Additionally, the methodology could be extended to incorporate more complex spatial relationships or to '
        'analyze temporal sequences of images.'
    )
    
    # Save the document
    doc.save(report_path)
    print(f"Report saved to {report_path}")
    return report_path

if __name__ == "__main__":
    report_path = create_gap_analysis_report()
    print(f"GAP Analysis Report generated successfully at: {report_path}")
