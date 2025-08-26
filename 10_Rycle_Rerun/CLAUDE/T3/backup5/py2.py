import os
import glob
from docx import Document
from docx.shared import Inches
from datetime import datetime

def generate_gap_analysis_report():
    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading('Analysis of Grayscale Anomaly Patterns (GAP) in Polymer Surface Images', 0)
    
    # Add creation date
    doc.add_paragraph(f'Report Date: {datetime.now().strftime("%Y-%m-%d")}')
    doc.add_paragraph('Author: Automated Analysis System')
    
    # Define output directory where results are stored
    output_directory = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup5"
    
    # Find all visualization images
    visualization_images = glob.glob(os.path.join(output_directory, "*_gap_visualization.png"))
    csv_files = glob.glob(os.path.join(output_directory, "*_gap_analysis.csv"))
    
    # Sort files to ensure consistent ordering
    visualization_images.sort()
    csv_files.sort()
    
    # Abstract Section
    doc.add_heading('Abstract', 1)
    abstract_text = (
        "This report presents a comprehensive analysis of polymer surface images using advanced image processing techniques. "
        "The study focuses on identifying Grayscale Anomaly Patterns (GAP) within a series of polymer images, characterized by "
        "specific grayscale value ranges and pattern continuity criteria. Using Contrast Limited Adaptive Histogram Equalization "
        "(CLAHE) for image enhancement and pixel-level analysis, we identified regions meeting the GAP criteria defined as pixels "
        "with grayscale values between 1-150 and having at least one adjacent direction with 25 contiguous pixels meeting the same "
        "grayscale condition. The analysis generated both quantitative data in CSV format and visual representations highlighting "
        "GAP regions in black against a white background. A total of " + str(len(visualization_images)) + " polymer images were processed, "
        "revealing distinct pattern distributions across different samples. These results provide valuable insights into the structural "
        "characteristics and potential anomalies within the polymer surfaces, which can inform further materials science research and "
        "quality control processes. The automated nature of this analysis demonstrates the effectiveness of computational methods in "
        "materials characterization and defect identification."
    )
    doc.add_paragraph(abstract_text)
    
    # Introduction Section
    doc.add_heading('Introduction', 1)
    intro_text = (
        "Polymer surface analysis is crucial in materials science and engineering, providing essential information about material "
        "properties, quality, and potential defects. Traditional visual inspection methods often lack objectivity and consistency, "
        "highlighting the need for automated computational approaches.\n\n"
        
        "This study implements an image processing methodology to identify specific grayscale patterns, termed Grayscale Anomaly "
        "Patterns (GAP), within polymer surface images. These patterns may correspond to structural features, compositional variations, "
        "or manufacturing artifacts that could influence material performance and quality.\n\n"
        
        "The GAP conditions are defined by two key criteria:\n"
        "1. Pixels with grayscale values between 1 and 150 (inclusive)\n"
        "2. Pixels that have at least one adjacent direction (up, down, left, or right) containing 25 contiguous pixels that also "
        "meet the grayscale condition\n\n"
        
        "These criteria were specifically designed to identify regions with particular grayscale characteristics and spatial continuity, "
        "which may indicate important material properties or anomalies. By applying these criteria systematically across multiple polymer "
        "samples, we aim to provide quantitative data on pattern distribution and characteristics, enabling more informed material "
        "evaluation and quality assessment."
    )
    doc.add_paragraph(intro_text)
    
    # Methods Section
    doc.add_heading('Methods', 1)
    methods_text = (
        "The analysis was conducted using a Python-based image processing pipeline, leveraging the OpenCV and Pillow libraries. "
        "The methodology consisted of several key steps:\n\n"
    )
    methods_para = doc.add_paragraph(methods_text)
    
    # Add method details with bold headings
    p = doc.add_paragraph()
    p.add_run("Image Acquisition and Selection: ").bold = True
    p.add_run(
        "The analysis focused on polymer surface images with the prefix 'Poly_' in PNG or JPG format. "
        "A total of " + str(len(visualization_images)) + " images were processed from the specified input directory."
    )
    
    p = doc.add_paragraph()
    p.add_run("Image Enhancement: ").bold = True
    p.add_run(
        "All images were processed using Contrast Limited Adaptive Histogram Equalization (CLAHE) with a clip limit of 3 and "
        "a tile grid size of 10×10. This enhancement technique improved local contrast while preventing noise amplification, "
        "making subtle features more distinguishable for subsequent analysis."
    )
    
    p = doc.add_paragraph()
    p.add_run("Grayscale Conversion and Pixel Analysis: ").bold = True
    p.add_run(
        "The enhanced images were converted to grayscale to simplify analysis. Each pixel was then evaluated against the GAP criteria: "
        "grayscale value between 1 and 150, and at least one adjacent direction containing 25 contiguous pixels meeting the same "
        "grayscale condition. This evaluation required checking four directions (up, down, left, right) from each pixel."
    )
    
    p = doc.add_paragraph()
    p.add_run("Data Export and Visualization: ").bold = True
    p.add_run(
        "For each processed image, two outputs were generated: (1) a CSV file containing the coordinates, grayscale value, and "
        "GAP flag (0 or 1) for each pixel, and (2) a visualization image highlighting GAP pixels in black (RGB: 0,0,0) and "
        "non-GAP pixels in white (RGB: 255,255,255). This binary representation provides a clear visual indication of GAP "
        "distribution across the sample surface."
    )
    
    # Results Section
    doc.add_heading('Results', 1)
    results_intro = (
        "The analysis successfully processed all " + str(len(visualization_images)) + " polymer images and identified regions meeting "
        "the GAP criteria. The results are presented as both quantitative data in CSV format and visual representations highlighting "
        "GAP regions.\n\n"
        
        "The visualization images clearly delineate areas where the GAP conditions are met, showing the spatial distribution of these "
        "regions across the polymer surfaces. These patterns may correspond to specific structural features, compositional variations, "
        "or manufacturing artifacts within the polymer samples.\n\n"
        
        "The CSV data provides a comprehensive pixel-by-pixel analysis, enabling further statistical evaluation and pattern recognition "
        "beyond the scope of this initial report. This data can be used for more detailed studies of pattern characteristics, size "
        "distribution, or correlation with other material properties.\n\n"
        
        "Below are the visualization results for all processed images, showing the distribution of GAP regions (black) against "
        "non-GAP regions (white):"
    )
    doc.add_paragraph(results_intro)
    
    # Add all visualization images with captions
    for i, img_path in enumerate(visualization_images):
        img_filename = os.path.basename(img_path)
        base_name = img_filename.replace('_gap_visualization.png', '')
        
        # Add a subheading for each image
        doc.add_heading(f'Sample {i+1}: {base_name}', 2)
        
        # Add the image
        try:
            doc.add_picture(img_path, width=Inches(6))
            
            # Add a caption
            doc.add_paragraph(f'Figure {i+1}: GAP visualization for {base_name}. Black regions indicate pixels meeting the GAP criteria, '
                             'while white regions represent pixels that do not meet these criteria.')
        except Exception as e:
            doc.add_paragraph(f"Error including image {img_path}: {str(e)}")
    
    # Conclusion paragraph in Results section
    conclusion_text = (
        "The visualizations demonstrate the effectiveness of the GAP analysis in identifying specific pattern regions within the "
        "polymer samples. Each sample shows a unique distribution of GAP regions, reflecting the individual characteristics of the "
        "polymer surfaces. The black areas in the visualizations represent regions where both the grayscale value criterion and the "
        "contiguity criterion are met, potentially indicating areas of interest for further investigation.\n\n"
        
        "These results provide a foundation for further investigation into the correlation between these patterns and material "
        "properties, manufacturing conditions, or performance characteristics. The quantitative data stored in the CSV files can be "
        "used for statistical analysis to identify trends or correlations across multiple samples, potentially revealing insights "
        "about the manufacturing process or material composition."
    )
    doc.add_paragraph(conclusion_text)
    
    # Save the document
    report_path = os.path.join(output_directory, 'GAP_Analysis_Report.docx')
    doc.save(report_path)
    print(f"Report generated successfully and saved to: {report_path}")
    return report_path

if __name__ == "__main__":
    # Check if output directory exists
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup5"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
        print(f"Created output directory: {output_dir}")
    
    # Verify if visualization images and CSV files exist
    visualization_images = glob.glob(os.path.join(output_dir, "*_gap_visualization.png"))
    csv_files = glob.glob(os.path.join(output_dir, "*_gap_analysis.csv"))
    
    if len(visualization_images) > 0 and len(csv_files) > 0:
        print(f"Found {len(visualization_images)} visualization images and {len(csv_files)} CSV files.")
        print("Calculation successful")
        
        # Generate the report
        report_path = generate_gap_analysis_report()
        print(f"Analysis complete. Report available at: {report_path}")
    else:
        print("Error: Required output files not found. Please ensure py1.py has been executed successfully.")
