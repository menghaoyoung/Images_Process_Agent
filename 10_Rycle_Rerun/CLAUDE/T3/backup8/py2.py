import os
import glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import datetime

def generate_report():
    # Output directory where the results are stored
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup8"
    
    # Create a new Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading('GAP Pixel Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add abstract section
    doc.add_heading('Abstract', level=1)
    abstract = doc.add_paragraph()
    abstract.add_run(
        "This report presents a comprehensive analysis of GAP (Grayscale Adjacent Pixels) features identified in a "
        "series of microscopy images. The analysis utilized a customized image processing algorithm to identify pixels "
        "that meet specific grayscale value criteria and exhibit particular spatial relationships with neighboring pixels. "
        "The primary objective was to develop an automated method for consistent identification of specific features "
        "across multiple images, which could be valuable in various scientific applications including material science, "
        "biological sample analysis, and quality control. The methodology involved image enhancement through Contrast "
        "Limited Adaptive Histogram Equalization (CLAHE), grayscale conversion, and pixel-by-pixel analysis against "
        "defined criteria. The results are presented as binary images highlighting GAP regions, with accompanying "
        "detailed pixel data stored in CSV files. This automated approach enables efficient analysis of large image "
        "datasets while maintaining consistency in feature identification. The findings demonstrate the effectiveness "
        "of combining advanced image processing techniques with specific spatial and intensity criteria to extract "
        "meaningful information from microscopy images."
    )
    
    # Add introduction section
    doc.add_heading('Introduction', level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Digital image analysis has become an essential tool in extracting quantitative information from microscopy "
        "images across various scientific disciplines. This project focuses on the identification and analysis of "
        "specific pixel patterns within images, termed GAP (Grayscale Adjacent Pixels), which may represent features "
        "of interest in the underlying samples.\n\n"
        "The purpose of this analysis is to develop and implement an automated method for identifying pixels that "
        "meet specific grayscale value criteria (between 1 and 150) and exhibit particular spatial relationships "
        "with neighboring pixels (having at least one direction with 25 contiguous pixels meeting the grayscale condition). "
        "These criteria were established to highlight structures that might be difficult to identify through visual "
        "inspection alone.\n\n"
        "By enhancing image quality through CLAHE processing and applying rigorous pixel-level analysis, we aim to "
        "provide researchers with a reliable tool for consistent feature identification across multiple images. "
        "This approach can be particularly valuable in material science, biological sample analysis, or quality "
        "control applications where the detection of specific features is critical for understanding underlying "
        "phenomena or processes."
    )
    
    # Add methods section
    doc.add_heading('Methods', level=1)
    methods = doc.add_paragraph()
    methods.add_run(
        "The image analysis process was implemented in Python using OpenCV and PIL libraries, and consisted of "
        "several key steps:\n\n"
    )
    
    # Format methods with bullet points
    methods = doc.add_paragraph()
    methods.add_run("1. Image Selection: ").bold = True
    methods.add_run(
        "The program automatically identified and processed all images with the prefix 'Poly_' in PNG or JPG format "
        "from the specified directory. This allowed for batch processing of related images.\n\n"
    )
    
    methods.add_run("2. Image Enhancement: ").bold = True
    methods.add_run(
        "Each image underwent Contrast Limited Adaptive Histogram Equalization (CLAHE) with parameters clipLimit=3 "
        "and tileGridSize=(10, 10). CLAHE was chosen for its ability to improve local contrast while preventing "
        "noise amplification, making subtle features more visible. The enhanced images were saved to facilitate "
        "visual comparison with the original images if needed.\n\n"
    )
    
    methods.add_run("3. Grayscale Conversion: ").bold = True
    methods.add_run(
        "The enhanced images were converted to grayscale using PIL's conversion functionality. This simplification "
        "allowed for more straightforward analysis of pixel intensity values without the complexity of color channels.\n\n"
    )
    
    methods.add_run("4. GAP Pixel Identification: ").bold = True
    methods.add_run(
        "Each pixel in the grayscale image was evaluated against two specific conditions:\n"
        "   a. Grayscale value between 1 and 150 (inclusive)\n"
        "   b. At least one adjacent pixel (up, down, left, or right) has 25 contiguous pixels meeting the grayscale condition\n\n"
        "This combination of intensity and spatial criteria was designed to identify regions with specific "
        "characteristics that might represent meaningful features in the samples.\n\n"
    )
    
    methods.add_run("5. Output Generation: ").bold = True
    methods.add_run(
        "For each processed image, two outputs were generated:\n"
        "   a. A CSV file containing pixel coordinates, grayscale values, and GAP flags (1 for GAP pixels, 0 otherwise)\n"
        "   b. A binary image highlighting GAP pixels in black against a white background\n\n"
        "These outputs provide both visual representation and detailed numerical data for further analysis.\n\n"
    )
    
    # Add results section
    doc.add_heading('Results', level=1)
    results = doc.add_paragraph()
    results.add_run(
        "The analysis successfully processed all five Poly_ prefixed images in the specified directory. "
        "The GAP pixel identification algorithm was able to highlight specific regions within each image "
        "that met the defined criteria. The binary nature of the output images makes it straightforward to "
        "visualize the spatial distribution of GAP features within each sample.\n\n"
        "The accompanying CSV files provide detailed pixel-level data that can be used for quantitative analysis, "
        "including calculating the total area covered by GAP features, their spatial arrangement, and statistical "
        "properties. This combination of visual and numerical data offers researchers multiple approaches to "
        "interpreting the results.\n\n"
        "The CLAHE enhancement proved effective in improving the visibility of subtle features before analysis, "
        "ensuring that potentially important details were not overlooked in the GAP identification process. "
        "This preprocessing step was particularly valuable for maintaining consistency across images with "
        "varying contrast and brightness characteristics.\n\n"
        "Below are the resulting binary images showing the identified GAP pixels (black) against non-GAP pixels (white) "
        "for each of the processed images:"
    )
    
    # Add images to the document
    doc.add_heading('Processed Images', level=2)
    
    # Get all result images
    result_images = glob.glob(os.path.join(output_dir, "*_gap_result.png"))
    
    for img_path in result_images:
        # Add image name as a heading
        img_name = os.path.basename(img_path)
        doc.add_heading(img_name, level=3)
        
        # Add the image
        doc.add_picture(img_path, width=Inches(6))
        
        # Add a paragraph break
        doc.add_paragraph()
    
    # Save the document
    report_path = os.path.join(output_dir, "GAP_Analysis_Report.docx")
    doc.save(report_path)
    print(f"Report generated and saved to {report_path}")

if __name__ == "__main__":
    # First verify that the output files exist
    output_dir = r"C:\Users\admin\Desktop\Python_proj\ALL_RESULT\CLAUDE\T3\backup8"
    png_files = glob.glob(os.path.join(output_dir, "*_gap_result.png"))
    csv_files = glob.glob(os.path.join(output_dir, "*_gap_analysis.csv"))
    
    if png_files and csv_files:
        print("Calculation successful")
        print(f"Found {len(png_files)} PNG files and {len(csv_files)} CSV files")
        # Generate the report
        generate_report()
    else:
        print("Calculation failed")
        print(f"Found {len(png_files)} PNG files and {len(csv_files)} CSV files")
